from __future__ import annotations

from pathlib import Path

from docx import Document


def _remove_toc_paragraphs(document: Document) -> None:
    """Delete paragraphs that belong to a table of contents."""
    for paragraph in list(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("TOC") or "......" in paragraph.text:
            element = paragraph._element
            element.getparent().remove(element)


def normalize_styles(doc_path: Path, out_path: Path) -> None:
    """Normalize visual styles to structured heading styles in a DOCX file."""
    document = Document(str(doc_path))
    _remove_toc_paragraphs(document)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            size = run.font.size
            bold = run.bold
            if bold and size and size.pt >= 14:
                paragraph.style = "Heading 1"
                break
            if bold and size and size.pt >= 12:
                paragraph.style = "Heading 2"
                break
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
