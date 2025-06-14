import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt
from typer.testing import CliRunner

from wiki_documental.cli import app

runner = CliRunner()


def _create_doc(path: Path) -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("Title")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Body")
    doc.save(path)


def _fake_run(cmd, capture_output=True, text=True):
    md = Path(cmd[-1])
    md.write_text("# Title\nBody", encoding="utf-8")
    class R:
        returncode = 0
        stderr = ""
    return R()


def _setup_paths(tmp_path):
    paths = {
        "originals": tmp_path / "orig",
        "work": tmp_path / "work",
        "wiki": tmp_path / "wiki",
        "tmp": tmp_path / "tmp",
    }
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)
    return paths


def _prepare_files(paths):
    doc_path = paths["originals"] / "sample.docx"
    _create_doc(doc_path)

    map_data = [{"level": 1, "title": "A", "slug": "a"}]
    index_data = [{"title": "B", "slug": "b", "children": []}]
    (paths["work"] / "map.yaml").write_text(
        yaml.safe_dump(map_data, allow_unicode=True), encoding="utf-8"
    )
    (paths["work"] / "index.yaml").write_text(
        yaml.safe_dump(index_data, allow_unicode=True), encoding="utf-8"
    )


def test_full_skip_verify(tmp_path, monkeypatch):
    paths = _setup_paths(tmp_path)
    _prepare_files(paths)

    monkeypatch.setattr("subprocess.run", _fake_run)
    monkeypatch.setattr("wiki_documental.processing.docx_to_md.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki_documental.cli.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki_documental.cli.cfg", {"paths": paths, "options": {"cutoff_similarity": 0.5}})

    result = runner.invoke(app, ["full", "--skip-verify"])
    assert result.exit_code == 0
    assert (paths["wiki"] / "_sidebar.md").exists()
    assert list(paths["wiki"].glob("*.md"))


def test_full_strict_verify(tmp_path, monkeypatch):
    paths = _setup_paths(tmp_path)
    _prepare_files(paths)

    monkeypatch.setattr("subprocess.run", _fake_run)
    monkeypatch.setattr("wiki_documental.processing.docx_to_md.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki_documental.cli.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki_documental.cli.cfg", {"paths": paths, "options": {"cutoff_similarity": 0.5}})

    result = runner.invoke(app, ["full"])
    assert result.exit_code == 1
