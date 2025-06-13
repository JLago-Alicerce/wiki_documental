import yaml
from docx import Document
from docx.shared import Pt
from pathlib import Path
from typer.testing import CliRunner

from wiki_documental.cli import app

runner = CliRunner()


def test_full_multiple_doc_sources(tmp_path, monkeypatch):
    paths = {
        "originals": tmp_path / "orig",
        "work": tmp_path / "work",
        "wiki": tmp_path / "wiki",
        "tmp": tmp_path / "tmp",
    }
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)

    doc_a = paths["originals"] / "DocA.docx"
    doc_b = paths["originals"] / "DocB.docx"

    for doc_path, text in [(doc_a, "A"), (doc_b, "B")]:
        doc = Document()
        run = doc.add_paragraph().add_run("Introducción")
        run.bold = True
        run.font.size = Pt(16)
        doc.add_paragraph(text)
        doc.save(doc_path)

    def fake_run(cmd, capture_output=True, text=True):
        out = Path(cmd[-1])
        out.write_text("# Introducción\nX", encoding="utf-8")
        class R:
            returncode = 0
            stderr = ""
        return R()

    monkeypatch.setattr("subprocess.run", fake_run)
    monkeypatch.setattr("wiki_documental.processing.docx_to_md.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki_documental.cli.ensure_pandoc", lambda: None)

    def fake_build_headings_map(_):
        return [{"level": 1, "title": "Introducción", "slug": "introduccion"}]

    monkeypatch.setattr(
        "wiki_documental.processing.headings_map.build_headings_map",
        fake_build_headings_map,
    )

    monkeypatch.setattr(
        "wiki_documental.cli.cfg",
        {"paths": paths, "options": {"cutoff_similarity": 0.5}},
    )

    result = runner.invoke(app, ["full"])
    assert result.exit_code == 0

    final = paths["wiki"] / "1_introduccion.md"
    assert final.exists()
    lines = final.read_text(encoding="utf-8").splitlines()
    assert lines[0] == "---"
    end = lines.index("---", 1)
    meta = yaml.safe_load("\n".join(lines[1:end]))
    assert sorted(meta["doc_source"]) == ["DocA.docx", "DocB.docx"]
    assert lines[end + 1].startswith("<div class=\"fragment-meta\"")
