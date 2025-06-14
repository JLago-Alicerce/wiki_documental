from pathlib import Path
from docx import Document
from docx.shared import Pt
from typer.testing import CliRunner

from wiki.cli import app

runner = CliRunner()

def _create_doc(path: Path) -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("Title")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Body")
    doc.save(path)

def _fake_run(cmd, capture_output=True, text=True):
    md = Path(cmd[-1])
    md.write_text("# Title\nBody", encoding="utf-8")
    for part in cmd:
        if part.startswith("--extract-media="):
            dest = Path(part.split("=", 1)[1]) / "media"
            dest.mkdir(parents=True, exist_ok=True)
            (dest / "img.png").write_text("bin", encoding="utf-8")
    class R:
        returncode = 0
        stderr = ""
    return R()

def test_reset_work_dir(tmp_path, monkeypatch):
    paths = {
        "originals": tmp_path / "orig",
        "work": tmp_path / "work",
        "wiki": tmp_path / "wiki",
        "tmp": tmp_path / "tmp",
    }
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)

    doc_file = paths["originals"] / "sample.docx"
    _create_doc(doc_file)

    monkeypatch.setattr("subprocess.run", _fake_run)
    monkeypatch.setattr("wiki.processing.docx_to_md.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki.cli.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki.cli.cfg", {"paths": paths, "options": {"cutoff_similarity": 0.5}})
    monkeypatch.setattr("wiki.config.cfg", {"paths": paths, "options": {"cutoff_similarity": 0.5}})

    result = runner.invoke(app, ["full"])
    assert result.exit_code == 0

    assert (paths["work"] / "md_raw").exists()
    assert (paths["wiki"] / "assets" / "media").exists()

    result = runner.invoke(app, ["reset"])
    assert result.exit_code == 0

    for d in ["md_raw", "normalized", "tmp", "media"]:
        assert not (paths["work"] / d).exists()
    assert not (paths["wiki"] / "assets" / "media").exists()

    for pattern in ["*.md", "*.yaml", "*.csv"]:
        assert not list(paths["work"].rglob(pattern))
        assert not list(paths["wiki"].rglob(pattern))

