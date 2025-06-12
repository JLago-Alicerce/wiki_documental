from docx import Document
from docx.shared import Pt

from wiki_documental.processing.normalize_docx import normalize_styles


def test_normalize_styles(tmp_path):
    sample = tmp_path / "sample.docx"
    doc = Document()

    run_h1 = doc.add_paragraph().add_run("Title H1")
    run_h1.bold = True
    run_h1.font.size = Pt(16)

    run_h2 = doc.add_paragraph().add_run("Title H2")
    run_h2.bold = True
    run_h2.font.size = Pt(12)

    run_h3 = doc.add_paragraph().add_run("Title H3")
    run_h3.bold = True
    run_h3.font.size = Pt(10)

    run_h4 = doc.add_paragraph().add_run("Title H4")
    run_h4.bold = True
    run_h4.italic = True

    doc.add_paragraph("Body text")
    doc.save(sample)

    out = tmp_path / "out.docx"
    normalize_styles(sample, out)

    doc = Document(out)
    assert doc.paragraphs[0].style.name == "Heading 1"
    assert doc.paragraphs[1].style.name == "Heading 2"
    assert doc.paragraphs[2].style.name == "Heading 3"
    assert doc.paragraphs[3].style.name == "Heading 4"
    assert doc.paragraphs[4].style.name == "Normal"
