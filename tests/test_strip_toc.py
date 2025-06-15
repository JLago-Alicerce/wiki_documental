from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from wiki.processing.normalize_docx import normalize_styles
from wiki.processing.docx_to_md import convert_docx_to_md


def _fake_pandoc_run(cmd, capture_output=True, text=True):
    docx = Path(cmd[1])
    md = Path(cmd[-1])
    doc = Document(docx)
    lines = []
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name == "Heading 1":
            lines.append(f"# {p.text}")
        elif name == "Heading 2":
            lines.append(f"## {p.text}")
        else:
            if p.text:
                lines.append(p.text)
    md.write_text("\n".join(lines), encoding="utf-8")
    class R:
        returncode = 0
        stderr = ""
    return R()


def test_toc_not_in_md(tmp_path, monkeypatch):
    docx_in = tmp_path / "in.docx"
    docx_out = tmp_path / "out.docx"
    md_file = tmp_path / "out.md"

    doc = Document()
    doc.add_paragraph("Contents", style="TOC Heading")
    doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Intro ...... 1", style="TOC 1")
    run = doc.add_paragraph().add_run("Chapter 1")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Body text")
    doc.save(docx_in)

    normalize_styles(docx_in, docx_out)
    monkeypatch.setattr("subprocess.run", _fake_pandoc_run)
    monkeypatch.setattr(
        "wiki.processing.docx_to_md.ensure_pandoc", lambda: None
    )
    convert_docx_to_md(docx_out, md_file)
    text = md_file.read_text(encoding="utf-8")
    assert "......" not in text

