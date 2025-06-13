from pathlib import Path
from docx import Document
from docx.shared import Pt
from typer.testing import CliRunner
from zipfile import ZipFile

from wiki_documental.cli import app

runner = CliRunner()


def _create_doc(path: Path) -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("Title")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Body")
    doc.save(path)


def _fake_run(cmd, capture_output=True, text=True):
    md = Path(cmd[-1])
    md.write_text("# Title\n![alt](media/img.png)", encoding="utf-8")
    media_dir = None
    for part in cmd:
        if part.startswith("--extract-media="):
            media_dir = Path(part.split("=", 1)[1])
    if media_dir is not None:
        (media_dir / "media").mkdir(parents=True, exist_ok=True)
        (media_dir / "media" / "img.png").write_text("binary", encoding="utf-8")
    class R:
        returncode = 0
        stderr = ""
    return R()


def test_package_static(tmp_path, monkeypatch):
    paths = {
        "originals": tmp_path / "orig",
        "work": tmp_path / "work",
        "wiki": tmp_path / "wiki",
        "tmp": tmp_path / "tmp",
    }
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)

    doc_path = paths["originals"] / "sample.docx"
    _create_doc(doc_path)

    monkeypatch.setattr("subprocess.run", _fake_run)
    monkeypatch.setattr("wiki_documental.processing.docx_to_md.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki_documental.cli.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki_documental.cli.cfg", {"paths": paths, "options": {"cutoff_similarity": 0.5}})

    monkeypatch.chdir(tmp_path)
    import sys
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

    result = runner.invoke(app, ["full"])
    assert result.exit_code == 0

    # Docsify landing page
    (paths["wiki"] / "index.html").write_text("index", encoding="utf-8")

    result = runner.invoke(app, ["package"])
    assert result.exit_code == 0

    dist = Path("dist")
    zips = list(dist.glob("wiki_documental_*.zip"))
    assert zips
    zip_file = zips[0]

    with ZipFile(zip_file) as z:
        names = z.namelist()
    assert "wiki/index.html" in names
    assert "wiki/_sidebar.md" in names
    assert any(n.startswith("wiki/assets/media/") for n in names)
