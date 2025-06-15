from pathlib import Path
from docx import Document
from docx.shared import Pt
from typer.testing import CliRunner

from wiki.cli import app

runner = CliRunner()


def test_pipeline_full(tmp_path, monkeypatch):
    paths = {
        "originals": tmp_path / "orig",
        "work": tmp_path / "work",
        "wiki": tmp_path / "wiki",
        "tmp": tmp_path / "tmp",
    }
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)

    doc_path = paths["originals"] / "sample.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("Title")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Body")
    doc.save(doc_path)

    def fake_run(cmd, capture_output=True, text=True):
        md = Path(cmd[-1])
        md.write_text("# Title\nBody", encoding="utf-8")
        class R:
            returncode = 0
            stderr = ""
        return R()

    monkeypatch.setattr("subprocess.run", fake_run)
    monkeypatch.setattr("wiki.processing.docx_to_md.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki.cli.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki.cli.cfg", {"paths": paths, "options": {"cutoff_similarity": 0.5}})

    result = runner.invoke(app, ["full"])
    assert result.exit_code == 0

    assert (paths["work"] / "index.yaml").exists()
    assert (paths["tmp"] / "tmp_full.md").exists()
    assert (paths["wiki"] / "_sidebar.md").exists()
    wiki_files = list(paths["wiki"].glob("*.md"))
    assert wiki_files


def test_pipeline_full_with_image(tmp_path, monkeypatch):
    paths = {
        "originals": tmp_path / "orig",
        "work": tmp_path / "work",
        "wiki": tmp_path / "wiki",
        "tmp": tmp_path / "tmp",
    }
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)

    doc_path = paths["originals"] / "img.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("Title")
    run.bold = True
    run.font.size = Pt(16)
    doc.save(doc_path)

    def fake_run(cmd, capture_output=True, text=True):
        md = Path(cmd[-1])
        md.write_text("# Title\n![alt](media\\img.png)", encoding="utf-8")
        media_dir = None
        for part in cmd:
            if part.startswith("--extract-media="):
                media_dir = Path(part.split("=", 1)[1])
        if media_dir is not None:
            (media_dir / "media").mkdir(parents=True, exist_ok=True)
            (media_dir / "media" / "img.png").write_text("binary", encoding="utf-8")
        class R:
            returncode = 0
            stderr = ""
        return R()

    monkeypatch.setattr("subprocess.run", fake_run)
    monkeypatch.setattr("wiki.processing.docx_to_md.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki.cli.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki.cli.cfg", {"paths": paths, "options": {"cutoff_similarity": 0.5}})

    result = runner.invoke(app, ["full"])
    assert result.exit_code == 0

    img_file = paths["wiki"] / "assets" / "media" / "img.png"
    assert img_file.exists()
    md_files = list(paths["wiki"].glob("*.md"))
    contents = [f.read_text(encoding="utf-8") for f in md_files]
    assert any("assets/media/img.png" in c for c in contents)
    assert all("assets/assets/media" not in c for c in contents)
    assert all("\\" not in c for c in contents)
