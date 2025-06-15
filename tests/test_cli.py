import builtins
from pathlib import Path
from docx import Document
from docx.shared import Pt

from typer.testing import CliRunner

import wiki
from wiki.cli import app

runner = CliRunner()


def test_version_option():
    result = runner.invoke(app, ["--version"])
    assert result.exit_code == 0
    assert wiki.__version__ in result.stdout


def test_full_calls_ensure_pandoc(monkeypatch, tmp_path):
    called = {"value": False}

    def dummy():
        called["value"] = True

    monkeypatch.setattr("wiki.cli.ensure_pandoc", dummy)

    paths = {
        "originals": tmp_path / "orig",
        "work": tmp_path / "work",
        "wiki": tmp_path / "wiki",
        "tmp": tmp_path / "tmp",
    }
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)

    doc_path = paths["originals"] / "sample.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("Title")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Text")
    doc.save(doc_path)

    def fake_run(cmd, capture_output=True, text=True):
        Path(cmd[-1]).write_text("# Title\nText", encoding="utf-8")
        class Result:
            returncode = 0
            stderr = ""
        return Result()

    monkeypatch.setattr("subprocess.run", fake_run)
    monkeypatch.setattr("wiki.processing.docx_to_md.ensure_pandoc", lambda: None)
    monkeypatch.setattr("wiki.cli.cfg", {"paths": paths, "options": {"cutoff_similarity": 0.5}})

    result = runner.invoke(app, ["full"])
    assert result.exit_code == 0
    assert called["value"]
    assert (paths["wiki"] / "_sidebar.md").exists()


def test_normalize_command(tmp_path, monkeypatch):
    sample = tmp_path / "sample.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("Title")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Text")
    doc.save(sample)

    monkeypatch.setattr(
        "wiki.cli.cfg", {"paths": {"work": tmp_path}}
    )
    result = runner.invoke(app, ["normalize", str(sample)])
    assert result.exit_code == 0
    assert (tmp_path / "normalized" / sample.name).exists()

import yaml


def test_map_command(tmp_path, monkeypatch):
    md_folder = tmp_path / "md_raw"
    md_folder.mkdir()
    (md_folder / "sample.md").write_text("# Title\n", encoding="utf-8")
    monkeypatch.setattr("wiki.cli.cfg", {"paths": {"work": tmp_path}})
    result = runner.invoke(app, ["map"])
    assert result.exit_code == 0
    map_file = tmp_path / "map.yaml"
    assert map_file.exists()
    data = yaml.safe_load(map_file.read_text(encoding="utf-8"))
    assert data[0]["slug"] == "title"
    assert data[0]["id"] == "1"


def test_index_overwrite(tmp_path, monkeypatch):
    map_data = [
        {"level": 1, "title": "A", "slug": "a"},
        {"level": 2, "title": "B", "slug": "b"},
        {"level": 3, "title": "C", "slug": "c"},
    ]
    work = tmp_path
    (work / "map.yaml").write_text(yaml.safe_dump(map_data, allow_unicode=True), encoding="utf-8")
    (work / "index.yaml").write_text("old", encoding="utf-8")
    monkeypatch.setattr("wiki.cli.cfg", {"paths": {"work": work}})

    result = runner.invoke(app, ["index", "--overwrite"])
    assert result.exit_code == 0
    data = yaml.safe_load((work / "index.yaml").read_text(encoding="utf-8"))
    assert data[0]["id"] == "1"
    assert not data[0]["children"][0]["children"]


def test_sidebar_command(tmp_path, monkeypatch):
    index = [
        {"id": "1", "title": "A", "slug": "a", "children": []}
    ]
    work = tmp_path
    (work / "index.yaml").write_text(yaml.safe_dump(index, allow_unicode=True), encoding="utf-8")
    paths = {"work": work, "wiki": work}
    monkeypatch.setattr("wiki.cli.cfg", {"paths": paths})
    (work / "README.md").write_text("intro", encoding="utf-8")

    result = runner.invoke(app, ["sidebar"])
    assert result.exit_code == 0
    sidebar = work / "_sidebar.md"
    assert sidebar.exists()
    content = sidebar.read_text(encoding="utf-8").splitlines()
    assert content == ["* [Inicio](README.md)", "* [A](1_a.md)"]


def test_sidebar_command_depth(tmp_path, monkeypatch):
    index = [
        {
            "id": "1",
            "title": "A",
            "slug": "a",
            "children": [
                {
                    "id": "1.1",
                    "title": "B",
                    "slug": "b",
                    "children": [],
                }
            ],
        }
    ]
    work = tmp_path
    (work / "index.yaml").write_text(yaml.safe_dump(index, allow_unicode=True), encoding="utf-8")
    paths = {"work": work, "wiki": work}
    monkeypatch.setattr("wiki.cli.cfg", {"paths": paths})
    (work / "README.md").write_text("intro", encoding="utf-8")

    result = runner.invoke(app, ["sidebar", "--depth", "2"])
    assert result.exit_code == 0
    sidebar = work / "_sidebar.md"
    content = sidebar.read_text(encoding="utf-8").splitlines()
    assert content[2] == "  * [B](1-1_b.md)"

