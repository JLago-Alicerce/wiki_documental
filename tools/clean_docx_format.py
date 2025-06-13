#!/usr/bin/env python
"""Clean and format DOCX files applying a DOTX template."""

from __future__ import annotations

import argparse
import logging
from datetime import datetime
from pathlib import Path
import shutil
import traceback

from docx import Document


def _remove_toc_paragraphs(document: Document) -> None:
    """Remove paragraphs that are part of a table of contents."""
    for paragraph in list(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("TOC") or "......" in paragraph.text:
            element = paragraph._element
            element.getparent().remove(element)


def _clean_paragraph(paragraph, template_styles) -> None:
    style_name = paragraph.style.name if paragraph.style else ""
    lower = style_name.lower()
    for i in range(1, 7):
        if f"heading {i}" in lower or f"título {i}" in lower or f"titulo {i}" in lower:
            paragraph.style = template_styles.get(f"Heading {i}", f"Heading {i}")
            break
    else:
        paragraph.style = template_styles.get("Normal", "Normal")
    for run in paragraph.runs:
        if getattr(run.font, "hidden", False):
            run.text = ""
        run.bold = None
        run.italic = None
        run.underline = None
        run.font.size = None
        run.font.name = None
        run.font.highlight_color = None


def _clean_tables(document: Document, template_styles) -> None:
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _clean_paragraph(paragraph, template_styles)


def clean_document(doc_path: Path, template_path: Path, dest_path: Path) -> None:
    document = Document(str(doc_path))
    template = Document(str(template_path))
    styles = {s.name: s.name for s in template.styles}
    _remove_toc_paragraphs(document)
    for paragraph in document.paragraphs:
        _clean_paragraph(paragraph, styles)
    _clean_tables(document, styles)
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dest_path))


def setup_logging(log_dir: Path) -> logging.Logger:
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / f"clean_docx_{datetime.now():%Y_%m_%d_%H_%M}.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
        handlers=[logging.FileHandler(log_file, encoding="utf-8"), logging.StreamHandler()],
    )
    return logging.getLogger(__name__)


def main() -> None:
    parser = argparse.ArgumentParser(description="Clean DOCX files using a DOTX template")
    parser.add_argument("--source", default="inputs/_originals/borrador", help="Input folder")
    parser.add_argument("--dest", default="inputs/_originals/listos_para_procesar", help="Output folder")
    parser.add_argument("--errors", default="inputs/_originals/errores", help="Error folder")
    parser.add_argument("--template", default="templates/plantilla.dotx", help="DOTX template")
    args = parser.parse_args()

    src_dir = Path(args.source)
    dest_dir = Path(args.dest)
    err_dir = Path(args.errors)
    template_path = Path(args.template)

    logger = setup_logging(Path("logs"))

    dest_dir.mkdir(parents=True, exist_ok=True)
    err_dir.mkdir(parents=True, exist_ok=True)

    for file in src_dir.iterdir():
        if not file.is_file() or file.suffix.lower() != ".docx":
            logger.info(f"Skipping {file.name}")
            continue
        logger.info(f"Processing {file.name}")
        try:
            out_file = dest_dir / file.name
            clean_document(file, template_path, out_file)
        except Exception:
            logger.error(f"Failed processing {file.name}\n{traceback.format_exc()}")
            shutil.move(str(file), err_dir / file.name)
        else:
            logger.info(f"Saved cleaned file to {out_file}")


if __name__ == "__main__":  # pragma: no cover - manual execution
    main()
